"""文档解析服务"""
from io import BytesIO
from pathlib import Path
from typing import Union

import pdfplumber
from docx import Document


def parse_text(file_path: Union[str, Path]) -> str:
    """解析文本文件

    Args:
        file_path: 文件路径

    Returns:
        提取的文本内容
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def parse_docx(file_path: Union[str, Path, bytes]) -> str:
    """解析 Word 文档

    Args:
        file_path: 文件路径或字节内容

    Returns:
        提取的文本内容
    """
    if isinstance(file_path, bytes):
        buf = BytesIO(file_path)
        doc = Document(buf)
    else:
        doc = Document(file_path)

    paragraphs = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            paragraphs.append(t)
    return "\n".join(paragraphs)


def parse_pdf(file_path: Union[str, Path, bytes]) -> str:
    """解析 PDF 文件

    Args:
        file_path: 文件路径或字节内容

    Returns:
        提取的文本内容
    """
    if isinstance(file_path, bytes):
        buf = BytesIO(file_path)
        pdf_source = buf
    else:
        pdf_source = file_path

    texts = []
    with pdfplumber.open(pdf_source) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            t = t.strip()
            if t:
                texts.append(t)
    return "\n".join(texts)


def parse_file(file_path: Union[str, Path]) -> str:
    """根据文件类型自动解析文件

    Args:
        file_path: 文件路径

    Returns:
        提取的文本内容

    Raises:
        ValueError: 不支持的文件类型
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == '.txt':
        return parse_text(path)
    elif ext in ['.doc', '.docx']:
        return parse_docx(path)
    elif ext == '.pdf':
        return parse_pdf(path)
    else:
        raise ValueError(f"不支持的文件类型: {ext}")
